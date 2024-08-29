import streamlit as st
import base64
from io import BytesIO
from datetime import datetime
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage
from PIL import Image as PILImage

AI71_BASE_URL = "https://api.ai71.ai/v1/"
AI71_API_KEY = os.getenv('AI71_API_KEY')

def get_llm():
    return ChatOpenAI(
        model="tiiuae/falcon-180B-chat",
        api_key=AI71_API_KEY,
        base_url=AI71_BASE_URL,
        streaming=True,
    )

def generate_resume_content(resume_data):
    llm = get_llm()
    
    prompt = f"""
    Generate a highly professional and ATS-optimized resume based on the following information:
    
    Name: {resume_data['name']}
    Email: {resume_data['email']}
    Phone: {resume_data['phone']}
    Location: {resume_data['location']}
    
    Work Experience:
    {json.dumps(resume_data['work_experience'], indent=2)}
    
    Education:
    {json.dumps(resume_data['education'], indent=2)}
    
    Skills: {', '.join(resume_data['skills'])}
    
    Please generate a compelling professional summary and enhance the job descriptions. 
    Use action verbs, quantify achievements where possible, and highlight key skills.
    Ensure the content is tailored for ATS optimization.
    The output should be in JSON format with the following structure:
    {{
        "summary": "Professional summary here",
        "work_experience": [
            {{
                "title": "Job title",
                "company": "Company name",
                "start_date": "Start date",
                "end_date": "End date",
                "description": "Enhanced job description with bullet points"
            }}
        ]
    }}
    """
    
    try:
        response = llm([HumanMessage(content=prompt)])
        enhanced_content = json.loads(response.content)
        
        resume_data['summary'] = enhanced_content['summary']
        resume_data['work_experience'] = enhanced_content['work_experience']
        
        return resume_data
    except Exception as e:
        st.error(f"An error occurred while generating AI content: {str(e)}")
        return resume_data

def create_docx(resume_data):
    doc = Document()
    
    # Styles
    styles = doc.styles
    style = styles.add_style('Name', 1)
    style.font.name = 'Calibri'
    style.font.size = Pt(24)
    style.font.bold = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add photo if provided
    if 'photo' in resume_data and resume_data['photo']:
        image_stream = BytesIO(resume_data['photo'])
        doc.add_picture(image_stream, width=Inches(2.0))
    
    # Add name
    doc.add_paragraph(resume_data['name'], style='Name')
    
    # Add contact information
    contact_info = doc.add_paragraph()
    contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info.add_run(f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}")
    
    # Add summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(resume_data['summary'])
    
    # Add work experience
    doc.add_heading('Work Experience', level=1)
    for job in resume_data['work_experience']:
        p = doc.add_paragraph(f"{job['title']} at {job['company']}", style='Heading 2')
        p.add_run(f"\n{job['start_date']} - {job['end_date']}")
        for bullet in job['description'].split('\n'):
            if bullet.strip():
                doc.add_paragraph(bullet.strip(), style='List Bullet')
    
    # Add education
    doc.add_heading('Education', level=1)
    for edu in resume_data['education']:
        p = doc.add_paragraph(f"{edu['degree']} in {edu['field']}", style='Heading 2')
        p.add_run(f"\n{edu['institution']}, {edu['graduation_date']}")
    
    # Add skills
    doc.add_heading('Skills', level=1)
    doc.add_paragraph(', '.join(resume_data['skills']))
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(resume_data):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
    
    story = []
    
    # Add photo if provided
    if 'photo' in resume_data and resume_data['photo']:
        image_stream = BytesIO(resume_data['photo'])
        img = Image(image_stream, width=100, height=100)
        story.append(img)
    
    # Add name
    story.append(Paragraph(resume_data['name'], styles['Title']))
    
    # Add contact information
    story.append(Paragraph(f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}", styles['Center']))
    story.append(Spacer(1, 12))
    
    # Add summary
    story.append(Paragraph('Professional Summary', styles['Heading1']))
    story.append(Paragraph(resume_data['summary'], styles['Justify']))
    story.append(Spacer(1, 12))
    
    # Add work experience
    story.append(Paragraph('Work Experience', styles['Heading1']))
    for job in resume_data['work_experience']:
        story.append(Paragraph(f"{job['title']} at {job['company']}", styles['Heading2']))
        story.append(Paragraph(f"{job['start_date']} - {job['end_date']}", styles['Normal']))
        for bullet in job['description'].split('\n'):
            if bullet.strip():
                story.append(Paragraph(f"• {bullet.strip()}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Add education
    story.append(Paragraph('Education', styles['Heading1']))
    for edu in resume_data['education']:
        story.append(Paragraph(f"{edu['degree']} in {edu['field']}", styles['Heading2']))
        story.append(Paragraph(f"{edu['institution']}, {edu['graduation_date']}", styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Add skills
    story.append(Paragraph('Skills', styles['Heading1']))
    story.append(Paragraph(', '.join(resume_data['skills']), styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_txt(resume_data):
    txt_content = f"{resume_data['name']}\n"
    txt_content += f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}\n\n"
    
    txt_content += "Professional Summary\n"
    txt_content += f"{resume_data['summary']}\n\n"
    
    txt_content += "Work Experience\n"
    for job in resume_data['work_experience']:
        txt_content += f"{job['title']} at {job['company']}\n"
        txt_content += f"{job['start_date']} - {job['end_date']}\n"
        for bullet in job['description'].split('\n'):
            if bullet.strip():
                txt_content += f"• {bullet.strip()}\n"
        txt_content += "\n"
    
    txt_content += "Education\n"
    for edu in resume_data['education']:
        txt_content += f"{edu['degree']} in {edu['field']}\n"
        txt_content += f"{edu['institution']}, {edu['graduation_date']}\n\n"
    
    txt_content += "Skills\n"
    txt_content += ', '.join(resume_data['skills'])
    
    return txt_content.encode()

def calculate_ats_score(resume_data):
    score = 0
    max_score = 100
    
    # Check for key sections
    if resume_data['name']: score += 5
    if resume_data['email']: score += 5
    if resume_data['phone']: score += 5
    if resume_data['location']: score += 5
    if resume_data['summary']: score += 10
    if resume_data['work_experience']: score += 20
    if resume_data['education']: score += 15
    if resume_data['skills']: score += 15
    
    # Check content quality
    if len(resume_data['summary'].split()) >= 50: score += 5
    if len(resume_data['work_experience']) >= 2: score += 5
    if len(resume_data['skills']) >= 5: score += 5
    
    # Check for keywords (this is a simplified version, in reality, you'd want to check against job-specific keywords)
    keywords = ['experience', 'skills', 'project', 'team', 'leadership', 'communication', 'achieved', 'improved', 'managed', 'developed']
    resume_text = ' '.join([str(value) for value in resume_data.values() if isinstance(value, str)])
    for keyword in keywords:
        if keyword in resume_text.lower():
            score += 1
    
    return min(score, max_score)

def main():
    st.set_page_config(page_title="AI-Enhanced Resume Builder", page_icon="📄", layout="wide")
    
    st.markdown("""
    <style>
    .big-font {
        font-size:30px !important;
        font-weight: bold;
    }
    .stButton>button {
        width: 100%;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Add sidebar
    st.sidebar.title("About This Project")
    st.sidebar.write("""
    Welcome to the AI-Enhanced Resume Builder!

    This project helps you create a professional, ATS-optimized resume with the power of AI. Here's what you can do:

    1. Input your personal information
    2. Add your work experience
    3. Include your education details
    4. List your skills
    5. Optionally upload a photo
    6. Generate AI-enhanced content
    7. Review and download your resume

    The AI will help improve your resume content and provide an ATS compatibility score.

    Get started by filling out the form and clicking 'Next' at each step!
    """)
    
    st.markdown('<p class="big-font">AI-Enhanced Resume Builder</p>', unsafe_allow_html=True)
    st.write("Create a professional, ATS-optimized resume with AI-powered content enhancement")
    
    # Initialize session state
    if 'step' not in st.session_state:
        st.session_state.step = 1
    
    if 'resume_data' not in st.session_state:
        st.session_state.resume_data = {
            'name': '', 'email': '', 'phone': '', 'location': '',
            'summary': '', 'work_experience': [], 'education': [], 'skills': [], 'photo': None
        }
    
    # Step 1: Personal Information
    if st.session_state.step == 1:
        st.subheader("Step 1: Personal Information")
        name = st.text_input("Full Name", st.session_state.resume_data['name'])
        email = st.text_input("Email", st.session_state.resume_data['email'])
        phone = st.text_input("Phone", st.session_state.resume_data['phone'])
        location = st.text_input("Location", st.session_state.resume_data['location'])
        
        photo_upload = st.file_uploader("Upload a photo (optional)", type=['jpg', 'jpeg', 'png'])
        if photo_upload:
            image = PILImage.open(photo_upload)
            st.image(image, caption='Uploaded Image', use_column_width=True)
            buffered = BytesIO()
            image.save(buffered, format="PNG")
            st.session_state.resume_data['photo'] = buffered.getvalue()
        
        if st.button("Next"):
            if name and email and phone and location:
                st.session_state.resume_data.update({
                    'name': name,
                    'email': email,
                    'phone': phone,
                    'location': location
                })
                st.session_state.step = 2
            else:
                st.error("Please fill in all required fields before proceeding.")
    
    # Step 2: Work Experience
    elif st.session_state.step == 2:
        st.subheader("Step 2: Work Experience")
        num_jobs = st.number_input("Number of jobs to add", min_value=1, max_value=10, value=len(st.session_state.resume_data['work_experience']) or 1)
        
        work_experience = []
        for i in range(num_jobs):
            st.write(f"Job {i+1}")
            job = {}
            job['title'] = st.text_input(f"Job Title {i+1}", st.session_state.resume_data['work_experience'][i]['title'] if i < len(st.session_state.resume_data['work_experience']) else '')
            job['company'] = st.text_input(f"Company {i+1}", st.session_state.resume_data['work_experience'][i]['company'] if i < len(st.session_state.resume_data['work_experience']) else '')
            job['start_date'] = st.date_input(f"Start Date {i+1}", value=datetime.strptime(st.session_state.resume_data['work_experience'][i]['start_date'] if i < len(st.session_state.resume_data['work_experience']) else '2020-01-01', '%Y-%m-%d')).strftime('%Y-%m-%d')
            job['end_date'] = st.date_input(f"End Date {i+1}", value=datetime.strptime(st.session_state.resume_data['work_experience'][i]['end_date'] if i < len(st.session_state.resume_data['work_experience']) else '2023-01-01', '%Y-%m-%d')).strftime('%Y-%m-%d')
            job['description'] = st.text_area(f"Job Description {i+1}", st.session_state.resume_data['work_experience'][i]['description'] if i < len(st.session_state.resume_data['work_experience']) else '', height=100)
            work_experience.append(job)
        
        col1, col2 = st.columns(2)
        if col1.button("Previous"):
            st.session_state.step = 1
        if col2.button("Next"):
            if all(job['title'] and job['company'] and job['description'] for job in work_experience):
                st.session_state.resume_data['work_experience'] = work_experience
                st.session_state.step = 3
            else:
                st.error("Please fill in all required fields for each job before proceeding.")
    
    # Step 3: Education
    elif st.session_state.step == 3:
        st.subheader("Step 3: Education")
        num_edu = st.number_input("Number of education entries", min_value=1, max_value=5, value=len(st.session_state.resume_data['education']) or 1)
        
        education = []
        for i in range(num_edu):
            st.write(f"Education {i+1}")
            edu = {}
            edu['degree'] = st.text_input(f"Degree {i+1}", st.session_state.resume_data['education'][i]['degree'] if i < len(st.session_state.resume_data['education']) else '')
            edu['field'] = st.text_input(f"Field of Study {i+1}", st.session_state.resume_data['education'][i]['field'] if i < len(st.session_state.resume_data['education']) else '')
            edu['institution'] = st.text_input(f"Institution {i+1}", st.session_state.resume_data['education'][i]['institution'] if i < len(st.session_state.resume_data['education']) else '')
            edu['graduation_date'] = st.date_input(f"Graduation Date {i+1}", value=datetime.strptime(st.session_state.resume_data['education'][i]['graduation_date'] if i < len(st.session_state.resume_data['education']) else '2023-01-01', '%Y-%m-%d')).strftime('%Y-%m-%d')
            education.append(edu)
        
        col1, col2 = st.columns(2)
        if col1.button("Previous"):
            st.session_state.step = 2
        if col2.button("Next"):
            if all(edu['degree'] and edu['field'] and edu['institution'] for edu in education):
                st.session_state.resume_data['education'] = education
                st.session_state.step = 4
            else:
                st.error("Please fill in all required fields for each education entry before proceeding.")
    
    # Step 4: Skills and Generation
    elif st.session_state.step == 4:
        st.subheader("Step 4: Skills and Resume Generation")
        skills_input = st.text_input("Skills (comma-separated)", ', '.join(st.session_state.resume_data['skills']))
        
        if st.button("Generate Resume"):
            if skills_input.strip():
                st.session_state.resume_data['skills'] = [skill.strip() for skill in skills_input.split(',') if skill.strip()]
                with st.spinner("Generating AI-enhanced resume content..."):
                    st.session_state.resume_data = generate_resume_content(st.session_state.resume_data)
                st.session_state.step = 5
                st.experimental_rerun()
            else:
                st.error("Please enter at least one skill before generating the resume.")
    
    # Step 5: Review and Download
    elif st.session_state.step == 5:
        st.subheader("Generated Resume")
        
        # Display resume content for review
        st.write("### Personal Information")
        st.write(f"**Name:** {st.session_state.resume_data['name']}")
        st.write(f"**Email:** {st.session_state.resume_data['email']}")
        st.write(f"**Phone:** {st.session_state.resume_data['phone']}")
        st.write(f"**Location:** {st.session_state.resume_data['location']}")
        
        if st.session_state.resume_data['photo']:
            st.image(st.session_state.resume_data['photo'], caption='Your Photo', width=200)
        
        st.write("### Professional Summary")
        st.write(st.session_state.resume_data['summary'])
        
        st.write("### Work Experience")
        for job in st.session_state.resume_data['work_experience']:
            st.write(f"**{job['title']} at {job['company']}**")
            st.write(f"{job['start_date']} - {job['end_date']}")
            st.write(job['description'])
        
        st.write("### Education")
        for edu in st.session_state.resume_data['education']:
            st.write(f"**{edu['degree']} in {edu['field']}**")
            st.write(f"{edu['institution']}, {edu['graduation_date']}")
        
        st.write("### Skills")
        st.write(', '.join(st.session_state.resume_data['skills']))
        
        # Calculate and display ATS score
        ats_score = calculate_ats_score(st.session_state.resume_data)
        st.write(f"### ATS Compatibility Score: {ats_score}%")
        
        # Download options
        st.write("### Download Options")
        col1, col2, col3 = st.columns(3)
        
        docx_buffer = create_docx(st.session_state.resume_data)
        col1.download_button(
            label="Download as DOCX",
            data=docx_buffer,
            file_name="resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        pdf_buffer = create_pdf(st.session_state.resume_data)
        col2.download_button(
            label="Download as PDF",
            data=pdf_buffer,
            file_name="resume.pdf",
            mime="application/pdf"
        )
        
        txt_content = create_txt(st.session_state.resume_data)
        col3.download_button(
            label="Download as TXT",
            data=txt_content,
            file_name="resume.txt",
            mime="text/plain"
        )
        
        if st.button("Edit Resume"):
            st.session_state.step = 1
        
        if st.button("Start Over"):
            st.session_state.step = 1
            st.session_state.resume_data = {
                'name': '', 'email': '', 'phone': '', 'location': '',
                'summary': '', 'work_experience': [], 'education': [], 'skills': [], 'photo': None
            }
            st.experimental_rerun()

if __name__ == "__main__":
    main()