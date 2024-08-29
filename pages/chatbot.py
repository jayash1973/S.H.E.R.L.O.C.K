import streamlit as st
import random
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
from langchain.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
import os
from dotenv import load_dotenv
import requests
from bs4 import BeautifulSoup
import pandas as pd
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import time
from langchain.schema import Document
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import io

# Load environment variables
load_dotenv()

AI71_BASE_URL = "https://api.ai71.ai/v1/"
AI71_API_KEY = "api71-api-92fc2ef9-9f3c-47e5-a019-18e257b04af2"

# Initialize session state variables
if "custom_personality" not in st.session_state:
    st.session_state.custom_personality = ""
if "messages" not in st.session_state:
    st.session_state.messages = []

# Initialize the Falcon model
@st.cache_resource
def get_llm():
    return ChatOpenAI(
        model="tiiuae/falcon-180B-chat",
        api_key=AI71_API_KEY,
        base_url=AI71_BASE_URL,
        streaming=True,
    )

# Initialize embeddings
@st.cache_resource
def get_embeddings():
    return HuggingFaceEmbeddings()

def process_documents(uploaded_files):
    documents = []
    for uploaded_file in uploaded_files:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        try:
            if file_extension in [".txt", ".md"]:
                content = uploaded_file.getvalue().decode("utf-8")
                documents.append(Document(page_content=content, metadata={"source": uploaded_file.name}))
            elif file_extension == ".docx":
                docx_file = io.BytesIO(uploaded_file.getvalue())
                doc = DocxDocument(docx_file)
                content = "\n".join([para.text for para in doc.paragraphs])
                documents.append(Document(page_content=content, metadata={"source": uploaded_file.name}))
            elif file_extension == ".pdf":
                pdf_file = io.BytesIO(uploaded_file.getvalue())
                pdf_reader = PdfReader(pdf_file)
                content = ""
                for page in pdf_reader.pages:
                    content += page.extract_text()
                documents.append(Document(page_content=content, metadata={"source": uploaded_file.name}))
            else:
                st.warning(f"Unsupported file type: {file_extension}")
        except Exception as e:
            st.error(f"Error processing file {uploaded_file.name}: {str(e)}")
    
    if not documents:
        return None
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    texts = text_splitter.split_documents(documents)
    
    vectorstore = FAISS.from_documents(texts, get_embeddings())
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    
    qa_chain = RetrievalQA.from_chain_type(
        llm=get_llm(),
        chain_type="stuff",
        retriever=retriever,
        return_source_documents=True,
    )
    
    return qa_chain

def get_chatbot_response(user_input, qa_chain=None, personality="default", web_search=False):
    system_message = get_personality_prompt(personality)
    
    web_info = ""
    if web_search:
        web_results = search_web_duckduckgo(user_input)
        web_info = "\n\n".join([f"Title: {result['title']}\nLink: {result['link']}\nSnippet: {result['snippet']}" for result in web_results])
        user_input += f"\n\nWeb search results:\n{web_info}"
    
    if qa_chain:
        result = qa_chain({"query": user_input})
        response = result['result']
        source_docs = result.get('source_documents', [])
    else:
        messages = [
            SystemMessage(content=system_message),
            HumanMessage(content=user_input)
        ]
        response = get_llm().invoke(messages).content
        source_docs = []
    
    return response, source_docs, web_results if web_search else None

def get_personality_prompt(personality):
    personalities = {
        "default": "You are a helpful assistant.",
        "sherlock": "You are Sherlock Holmes, the world's greatest detective. Respond with keen observation and deductive reasoning.",
        "yoda": "Wise and cryptic, you are. Like Yoda from Star Wars, speak you must.",
        "shakespeare": "Thou art the Bard himself. In iambic pentameter, respond with eloquence and poetic flair.",
        "custom": st.session_state.custom_personality
    }
    return personalities.get(personality, personalities["default"])

def search_web_duckduckgo(query: str, num_results: int = 3, max_retries: int = 3):
    api_key = os.getenv('api_key')
    cse_id = os.getenv('cse_id')
    
    for attempt in range(max_retries):
        try:
            service = build("customsearch", "v1", developerKey=api_key)
            res = service.cse().list(q=query, cx=cse_id, num=num_results).execute()
            results = []
            if "items" in res:
                for item in res["items"]:
                    result = {
                        "title": item["title"],
                        "link": item["link"],
                        "snippet": item.get("snippet", "")
                    }
                    results.append(result)
            return results
        except HttpError as e:
            print(f"HTTP error occurred: {e}. Attempt {attempt + 1} of {max_retries}")
        except Exception as e:
            print(f"An unexpected error occurred: {e}. Attempt {attempt + 1} of {max_retries}")
        time.sleep(2 ** attempt)
    print("Max retries reached. No results found.")
    return []

def main():
    st.set_page_config(page_title="S.H.E.R.L.O.C.K. Chatbot", page_icon="🕵️", layout="wide")
    
    st.title("S.H.E.R.L.O.C.K. Chatbot")
    
    # Sidebar
    with st.sidebar:
        st.image("data:image/jpeg;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/2wCEAAkGBw8QEBUQEBAVFRUVFRYXFhcVFhUVFRUXFRUWFhUYGBYYHSggGBolGxcVITEhJSkrLi4uFx8zODYtNygtLisBCgoKDg0OGhAQGismHSUtLS0tLS0tLS0tLjArLS0rLS8vLS0tLS0tLS0rLS0tLS0tLS0tLS0tLS0tKy0tLS0tK//AABEIALcBEwMBIgACEQEDEQH/xAAcAAABBQEBAQAAAAAAAAAAAAAAAQIDBAUGBwj/xABAEAABAwIDBQUFBgUDBAMAAAABAAIRAyEEEjEFQVFhcQYTIoGRMkKhscEHFCNSctEzYoKy8JKi4VNjc8IVJJP/xAAZAQADAQEBAAAAAAAAAAAAAAAAAQIDBAX/xAAtEQACAgEDBAEBBwUAAAAAAAAAAQIRAxIhMQRBUWETIjJxsdHh8PEFFCNikf/aAAwDAQACEQMRAD8A8gQhKukxESoQgAhCEIAEqEIAEJWtJsASeQlTswNZ2lJ5/pd+ydCK6FtN2EXYapWa4l1KC5uWBk3uDp1GpHAFYyGqBOwSpE5jCTABJOgFyegSGIpaFBzzDRPHcBzJNgOZVqjgmtvUMn8rSP8Ac7QdBJ/SVcD7AAAAaNFgOfM8zJ5qlGwI8PgqbLu8bv8AYPLV3nA5FXZJ1P8AwOAG4clCwKZqtKhFTbf8Ol+ur/bQUWxAczjBsBPK8SeCl23/AA6f66v9tBQ7EfleYMGLRrz+aUfthLg6RgcDTIN4sQYn8R8gHjfRPrFjnOD/AAGXXAtr7zd3VvoUwkHKCCDFyBA/iP8AaaNOo9EzE17vDhmEujiNdDw5GQt2ZRZo43EVGueyqGvGZ4aTfKcxnI8GQAfdmOIlUGlS42oW1aoa6QXvm387tQbTwKgaUkqNNVlzD0/ERIEB+vJrlbwLGBxFWm5wyOMsuWw10OEGHNBgnprrNGm8ZydxD/i1wHzVrZ20H0ScsEOa5pafZOZpb9fgm7rYEMmD4SbGx0PI8lZD2lozC8nxDXRuo37+CpzwUoJyjhJ9YbP09UxovtYW0ybFvHxEe0BedDY6QdVEGg6GDwP0P7/FSYKtkBdE2IIvcFzdZtv3cLqSq2i4ZmHKd4P7ft6JWVRBBFjZT92YkXHEfXgoXNI18t4PQpzDCLGkStUrUtE03WdLTucLt82/UehTqlFzLm4Ojhdp6H6apat6L09wQkQmFHjaVCFyGQIQhABCVCEAC0djbIfiXZWkNAjM4zYHfA1jVZy2djYeu2S0hoc286mQQCBusT6qoq2TK62NRuFZgnGn7VVpuYsOBAOtuNr2B1TnY+o7V7o4ZiqAoYh9YUiHPquIgTmc6RIud0eg4Qu52Ps7AYamKlUNr1MxaXOP4Ac0AuFNp/iBtvFGvDRaoznl0r36MXZfZuvjmltKi8tMkvaQweriGv8A0/LVGH7H0/x8PiWihVpsBY4ugEn2XSf4jD1t106LbHas1WdzRc5skZiyWANbcgEXvCXbG2D/APG0quIcC+niCxrnQahb3eaAdZ0vyBOkoa7syU5v0eUtwLmn8XwRu1eejfqYHCVYa8NEMGUHXe4/qdw5CByUdaqXvc86ucXHq4yfmnNWaR1WSMUzVGwKVqoRKwKVoUbFK1MaKe2x+HS/XV/toLJaYuCtfbn8Ol+ur/bQWQFk+SmauD2w5uUPkhuhGovP1PqtB1YPzOabGT6rm1JRquaZaY+R6jetY5PJDh4O2rmi+pUD5pu7x8OElh8Rs5uo6iddFTIgkawYtp5cll4XaTXWfY8dx893mtFpWip8EtlqlUEQ5sjiLOH7+flCk7q0tOYctR1G7rpzU+DxDe67upSm1QsdcXDS4j4buUqCm2TNMmeHveRHteV+SYJ2I0qxSrkDLYtmYOkmB1Gg0T/vFOp/Fbld+dgsT/OzTzEHkVHVolu8EHQtMgwAeu8agJJ9maUX2NDqZy21sTr4mzG7hwmRqq2inZkdTPumNY8J8Q4S718oCWjUyWqMD2kcbxxY8f8AI5JXRaQ2lUI6bxuPkpm5TpY8Dp5H9/VSfcQ8TQdn4sNqo8vf6t9AqrUKSfBWlonClaVA0qRpTGTITZQmB4+hCFyGIIQhAChCEIAVuq9CxmJod2DT91rHQTcHMGOETYEOH+gLzwCVpsxzw3K6DoOcDSdyuNCZ1+1K7WYem9kCpXz03PYIIpUzJY3SJc+535YWNTpHc4H4H0KtbaxDn08Jmv8A/Xzk8O8qOj4NCrUCtkc8FSLeEYRcq12sBGz8O25L69SpxMNYGT8UuzcI+q8NG8rfoPo4l7qTxk7lwpsdGWtSgjJUEkh7S4THNog6qcstMS4Q1zS/fo8uYpWhd52own3nMHUWDEsJZmpCDVe0xBE+IOEEbxmGtweHdTc0lrgWkGCCCCCNQQbgqS1fdUxzQpmqNgUzAgZI1SsUbQp6bCdEykZ+3v4dL9dX+2isdae2sSxwYxjg7KXkkaS8MEA74yai11lrJ8jY9ASBKgQqtYXGvp2FxwP04KoHDj/g1TgqTobVnT7M2oCTlNyCC12hlpGmjomVeo0w8wC1hO5xhun5jp5+q4sLSwe1XNs/xDj7w/daqfkjTR02JbUactQEOHEXI3X94cDfkm5zAG4En1gH5BJgdp5mZQRUZ+V3uzw3sPTXmiLTFiSBv0gx6EKlfce3Yu067C2HAydXCN1x4bT6gqUVHNDdHNAji10ucbg6HxEcbWWe1WWVyMpbYtaR1lznekOiEmi4stjI67Dkd+Um39L93Q+pT8TiXvgVLuHvEQ86RmO/zVY1Gu3ZTy9k+XunpbkE6nVtBuOB3dDu+SlLua2PaVK0puGoOqEhjSTBMb4HAbz0QFVhRNKE0ITGeRoQhcpzghCEACsChEF5yzcD3iOMbhzPlKrO0PQq9jzNap/5H/3FADAdzRHxJ8966HC9kMQQHPygb2gy+N40gHzWNgKdMmakkcGmCf6oMLrNm4DuMmMwZe6jmLcRScZLbEkzvsCRvkRvhRk1pfSb4Fjb+v8Aj37MXaOMqMLWvpAtYA0NdbIBoGPbDvUuE81s7MGCfT711Z1MAw5rruDomGua2X24CeS3ttbOwjnsrWqUz7TQSDEeLS+niHNsb1xlfDCmXU2OzNDjB3OvE9SAFfT5vkjqX/DPqun+LJp7c/ebGP7RNYzJg6bmXBNV1nGCD4W3MGNSdJEBaeOxQxBw+LaMhewsqnRssAIcTuA4ncBwWRtDGfeG02Npgd2HXi5zRYngI+JVHbtYNw1HDbw51Q/2j5u9FtJOjGD3TaOlr7Vp16uZr5ZNiBvDWtJte+UKDav3TFuHeOLKosHt8bnjc17d8bjIgW4Rx2EqugsBIvI6xBHmPkFcwFnibSD8kopaVFGkpXLWyGpRLHFpixi2h5hPYFLtHK12ZxABA84GWw1JssqvtI6U5aOPvn09nyvzSbolGjXxDKXtm/5R7Xnub534ArKxe0H1PD7LfyjTlmOrj18gFUQobsoEIQkIexpJAAJJIAAEkkmAAN5JXY9lKOz8K9lXaNM1ZuKYAcGc3MNnc58tFzew8R3ddr5gjNB4FzHNB+PkYKKtYueSTefluRLZG2KKfJ79gcIzaLs/eYPEbPygMo9we8Y7+aXeAi27yGq4r7QPsxZSpvxWzw6GAuqUCS4hupfTcbkC5ymbaaZVyvY/tJUwGIbVa7wEgVG7nNm8jeRqP+V6Jt6pWrPx2z6tV1XNQbi8G7wtOVviNMFoEgwW3kwLyslJmzxpnixw7w3MWmOMcePBMC3KVYHW8gg8wRBWIWkWIWmOerkyy4tFUPpVHNMtJB5Lawm2pAZUtBJzDiQ0GRu9kLDCVaptGNHY0HtMGZHIj4K45hIaWkOAbu1AzOPibqNddNLri8Lin0z4T1BuD5LdwG1KbyA7wnhMT0ctFKxo1TG74pzSrFBjC3xaS3xNALwPFMtm/wAOJTq+Ac1udpFSn+dlwJ/MNWHr8UOSTpmkVa2HYJjXGDUyH3SQcs8yLt6wVo1qhkNxTCSdKjYzEccw8NQdfULGaruFxr2DLZzDqx12ny3HmIKzlF8msWWfuYPs1qZG6SWnzaRZCYX0DeKjeQLSB0JukRcv2h7HjqEIWZxglSJUANfoehV3Hfxan/kf/cVTdoei1amGc6o9+U5e8eZ4+I258EDoTZ7SINwdea6vYWLM1mmwdSc525uendro059SszBYVpjPb5qvR2w/vRTYQKZe1ptJc0uuCTo2STAhXlVQpBha+RNnUbPk0abSIlrBHCQBCzcW2m94c2YLWSIiHZQHeUhLj9q5PCyQ+SNJiLE9Vm0sQ8kC3kCD6FYdJFx1OS5Ovr5wk4xi7pGrSgQNJMW5rB2/m+8vDhEQAL+yAMuvEX810WzsBWqPGUXF76N4F3ASqfbygKL6Fy6KLWZtM5YTMn+oLaeaOpROddPLQ59kc9TB6c9I89y7TbG2cPVouyvEio0Nc5pgSYLgNSMua68+fWLiJ0BmBop3Yy0AW+vPiOSiUIyab7BjyyhGUVwzR7XiiMRlovLgGAOJ4y4gC2mUhYiVziTJMk6pFTdsyWyBCEJDBCVCAA6LvcVsOlXY0s8JDQA4DcBoRvC4OF6v2YoMxeEbA9psGGg+NvhcYcIJzCbhZ5ZaVZ29HFSbTOB2ls12Ge0PMgiQYgW1Hlb1Xp+Bc/ucFjHgitTodyzK5ozMnLmfnEG8tDd9+S08f2Sw1TCFlTMDl9wCZF2hrQLutoF5ue39UYdtBuHbmptNNlR7iSGhxLc1GI7xsm8xyO8xNPcfUfQ9jLxlFtGtUogginUewdGuI+iVjwdQCN4NwscV3akyTckkkknUk71Yw2JF80RHqplifIQzR4ZW6JU1OC6TjHBOCYE8IGaGB2pUp2nM21ibiNIK6jZW2g45mPIdviz435m6PH+SuICewkGQYI3jVVfZlLY77E1g90hjW2vlkNJn2g33egskYCdFzWA20R4aokfmGvmN66PB4sEZmOBHI6fUFNV2NE7JJQmOeJSJlnlqEqRYHGCVCVADXaHousxUsc4OGh063HzC5N2i7ftGAWUakXLIPOGtI+ZSbVqLXJpCMqcovdGPicSDTcGyDF40IsCJ81FsBrPvDDUALQ5sg6eIhonoTPkoSPC7p/7NUeHNiNMwgdQQf3VuNqiIyqWo7TF0MPVcWYgFlRpguBgui0mZBBtqPNPoYLA4cZ34gvG5oIvHJpJPlC5urtbEFw8QLratDiToLEEHpC6HEPqVqXc1WsbiWtFVvdtDJDblj2i2fKMwjpvXJHFljs3t6O+fUYJO4xWr3+hNS2u51QMYCymXtBFgQHHKNNLkc+ay8ZjS+pXe6o0jvC0UXNLmva05egED16qPA4c942o9w8Lg4+RBsqW0fDiKo/7j/wDc6VvGMI5KXg58ksk8Ny4bMjaWHFOq5rfZsW/ocA5nnlcPOVWWptttqc65XD+mczZHV7vRWOzPZTGbQflw9Pwgw6o61NvnvPIfBW1TOW0lbMNdr2Q+zbG46KlQGhR/M8eNw/lZ9THQrusJ2J2fsei3EYhv3iuTDc3hbmAk5W6NA43O6brrOzfaU1srHUWAF+QZDGURI8J13aHidyVOrROpvjY4P7SeyGC2dskDD0hn76mDVdDqh1nxbuggX0XkC+gftpw1SrgWUqTC9zsRThouT4XH05rzCj2No0mj77iSyoRIZSGctG7MYv8ADqU4pyQ47bM45C0dq7LNBxyu7xm54Bb/AKmn2T6jms9FGhYwGFdWqspN1e4DoN58hJ8l7Zs+h3OHIoCCxoFMa2br1JuvPexOyHtecRVaW2hgNic3tOjda3mV3wxmUQNfkFksfyzt8L8f0O7HP4If7S/BfmXNgYlz64qYioSQCGTZocbWAsLTdQ/aL2DoVI2gykM7ROIpguaKo/6hyEHM3eQRIudL6XZnEMFSHMbJu10CQRrdab8VnqOJvci/DSPRYTl8DTZvOH9xxtR4ljOy7Kl8I7K7/o1HDxcqdUxf+V/H2iuZr0X03Fj2lrmmC1wIIPMHRel7bwIoYh9NvszLf0uEgeWnks7a2FbiaeWpdzR4H+83kT7zeR8oXeqkrR50sTicGnBNIIJB1Fj5JwSIFCcE0JwQMcE4JgTgmMepsPiH0zmY4g/PqN6hSoKRfqbWrOM5gOQAj4oVFCY7MZCAhZGAqEITAQhdhjsfSr4djZh1MgERNvZJHHpyXIJ/eO4nyt8lLim0/BcZuKaXdUaeNqimzuxlLjcuAeCBaxzRa3DeswklIhW3Zmkarmk02VGm8C+8OBj5iVNhcW8PFTMe8Ds2bU5pmTxuoNl1vw3s841kGM0DkWtKfs5ryT3YvxsCJMWJ39E07HJad0aON2mW1nN7toBALBMRIktPMSR/SqVA1alUugtJkkmQGiNSdR+8cVQrkFxN+AnW1vnKdSruFs0DfrB4SBqohjjB2jTJmnkjpkw2uCKrmH3LTx3z8V7Z9h9UHZrm/lrP+IB+q8KqmSTxXtP2I1IwVUf94/Fo/ZU1dnPk+ybvavDYipi2OcD3dIEsOWWw5rRM7zmBPLwrR7N4QMcaz2gADw7ogRm9LTzTdudscNhAWF+d4H8NpEj9btGDrfkVwWMr7T2qZfUFDCn2YDgHji1vtVt17M5hCtxolRbeo3O2/bPDktZRIqvY8uJB8A8Lmxm3+1u4Lg8dtB+IfneGzEDKItc+eq7PZ2yMNhRDKIeSCHOqwXvabOAOlMXjw3Fpc7Rc1t3Ypw7u8phzqDz4HEeyde7fweINt8SLLWDSVG0YrkpUKRfaJ48PNSYPZGFov7zKC4XAuQ39Ldx/wQjD4hwhrRr6k8eanLHtPjaQdYcCDffBWU4uWz4O3HpW/LKO0e2FJvhpMc4yM2YFkAG9jfN1AW5g8SHw4GQ6CDxB0WRtHY9LEDxeF8WeNeh/MP8ALKLCUqmGptZUIMSARoRJLfhFuSrBFQWlGXUvJKWtnc7Pc4OB4ELaNTxSN91weC208awfgfVdLsfG96HcQRbhI/4XJ/UYf49Xg7ugyXPT5M3tgPx2O40x8HOWFiHQx3T5kD6rb7Vn8Vn6P/YrnNp1MtInmPhf6LTpneGJn1Cqcji6zpc48XE+pKaEgSrc84cnBNCUJFDk4JoTgmMcnBNShBSFQhCBmOhCVQYAhCEACVIlCAFCUJE6nqOo+aYEzPDpqN+9WHY18RmO/wCO88SnVcLq4G2sR4tdI9b6WVZtJzvdMb+QS1IbhIkpUczZm4PrPyP7ow+Eq1H93Tpuc7gB89wHM2Wv2d2KMQ8uc/K1sSBGZ3rYC3xXfYLC06TctNoaOA/y61jik/uMpZopV3OFGwW4dzRiG5nOEhod+HwMkQSRbfF960MftmtggcLhH9zTqhj3EAlwJBBDXXIbbdfUSrvarCB1am4uc05CAWxuM6EEbxuVgdmKeJw7cry2o0uOZ3izzlHi0IENERESbKHikp7cGyyQliSa3NfZHZXCYXLVcRiXvaHtqvymkQ6HB1KmZDt3jfmvuatl9V1R8NzPe7hLnH14c7DeYIWH2VoYvDNOFxdJ1ShJdTqUXMc+i/UhofEsdvBEA34ldfs6g6uwhpGHoSQadF+au8ixFauLtP8AKwyPzRZD+lbhDHrexVpUWtf3b2mvWEfgUTanpBrVjanbdMkEjxhag2N3mU41zXhplmHpiKDDe5BvVdc3dA4NCv4SgykwU6LAxo3NEa6nrzRiK9OkAajoJ9kCS5x4NaLuPRZObZ1LEol2iWkAQBFgBuA+QWftXYuFxbvxRmewAS1xDgDJAdG7WJ5podWrazRZ+VpHeu/U4WZ0bJ/mCuUxSoM91jRfqTqeJPxKXBSKGH7HYBpnui79T3n4AwrWOpYDD0nMqUqQY4QWZW+Icx9SuP7QfadQY/uMLNR+hc0Ahsa3Nvn5Lkm4ivjcZTDqpGZtUgm4GXLfLOu6Sq07apOkiG23pW7IsXhB94LMO0kOLixupa3MYBPIRcrrNi7NNBpzOlzomNBEwBx11VepTdhz3eHYyY8dSoSST0Fz6gCVibVxeMa/LUrWIkBgyNPpfyJWOeObqI/Sqh75Z0dPLD07qTufrhFntZtCm0ucbim1oMcS68f6lyu18YypQmm4EX01HgdqNRqtB2DbWYWPmDwMGQZBXOY/ZL8NmkhzS05SNehbu+S6IRUIqC7HPllJty7MykqQJVRyihOCaE4JlIcEoSBKEDHBKE1OCBoVKkQgoyEqEKDAEqEIoYJYTSUkoESMbJhW6dIDcqmHMOCugqomc2xXYl7RlOm4/sdx6Kdm0R3ZDhLoI5Hmog7dqOB0TO5Z/N0kfNKeJT5NMfUOF0TYPHvpnOwwRcjcRMfVdXhu0L6gDB4XcYmf2/zRchU9kgCBbTrvO9TsJFwtk3wc7rmjrq2y8XXLZzSDYvPhE6/TRdXsTBOo08r3BziZJFhusFxew+0r6UMqeJvxC7zAYunVbmpuBHxHVWkGtvYvU08UXB/e0nZKm8xLXgaCo33hwNiNxF5bTVliiRtB09ib7/iXjK2m2kfeeSKn/wCYET1dH6Sp8JhWsJfcuPtPeZeerjoOQgDcFnbT2vQwlPvK9QMG6dXHg0auPILzHtN9otfEE08NNKnPtH23fRvxPRYOKXB1fJe8mejdpu3WEwIy5s9Tcxtz58BzMDqvO8f2qxGPGZ7srTIyA2HInfu4DkuFcSSSSSSZJNySdSTvK1NhVLuZxgjysfp6KscUnuRPK2qWyL+ApAVKvHOB5ZQfmStvYlUtxVFw1FOp8cqwGPLMU4HSoAR1A/4cFbxYeKfes93K13IPdr6tA81pSdpkptK1ydnidpU6cl7pdrA18+CoY/O9or1w2jSE5XVDlDp/KDd56BcXS7Q1aZmk1rXDRxAcRzANgesrPxeNq1nmpWqOqPIgueS4xMwJ0HIWTlO9kZxWl2ufJvbQ7SMHhw7Sf53jL/pZr6+iwziH1M7nuJOU69dw3KtCkpe9+krNsttvkjCVIE5AgCcE0JwQMUJQkCUIGOCVNCcgYqEiEx2ZaEIUGYqQlCQoYAhCEhApqVchQoQI0GPBTws5riFapV+KtSM5Q8E9X2T/AJvUrVDUcMuvD5hSGq0a+itE02iw1WMFtx+FdNN0neBp/nJY9bEudYWHD91Chz8FrGu565sXtpharR3jgx2+SA0czJsPhzVXtF9o9GlNPAt71+hrPBFNvHIzV55mBvuvLv8APiEilystRLO0MfWxFQ1a9R1R53uPwA0aOQsq8ISqShWMJIABJNgBqSul2PssU35arHh5Eh4g0wBctzDQ9fLnibKxncVRUyh0SI013g7itDavaB9ZpYxuRp1vLnDhyCuOlbsiV3SL9emyvUbXh4w9EwXsaXFxkE5QPcEXd1hS18c2lQqM1NQBreGrpd5W+CbsztcyjQbT7klzGgC4DDG8nUTroudqYh1R+ZxuZPIW0A3BU2uUON00ysE4BKGp4asyhoansAuOR+RShieGoAgyJIVktTSxFDIEqeWJsIARKkSoGKnJqVMBUJEIGZqEIUGYFIlQkAiEqEACEoanAIoKEa1PQhMYqVNSpgKhCUIAB9PqEJR9PqEIEASoSoASEsJQE8NTAYApGC48/iE5rU8NQMaGp4anAJYTAQBOhEJYQMSEQnQiExjCEwsU0JIQIrlibCswmliVDK6VPLE2EACRLCRAzOlCEKDMEIQkgFhKAhCoBUqEIGCVCEDFSoQgTBKhCYhR9PqEqEIAUJwCEIAkaE4BCEAPATgEITGOShCExipQhCBioQhAhEQhCYwhJCEIAQhNLUISAbkQhCBn/9k=", use_column_width=True)
        
        st.subheader("📁 Document Upload")
        uploaded_files = st.file_uploader("Upload documents", type=["txt", "md", "docx", "pdf"], accept_multiple_files=True)
        
        st.subheader("🎭 Chatbot Personality")
        personality = st.selectbox("Choose chatbot personality", ["default", "sherlock", "yoda", "shakespeare", "custom"])
        
        if personality == "custom":
            st.session_state.custom_personality = st.text_area("Enter custom personality details:", value=st.session_state.custom_personality)
        
        st.subheader("🌐 Web Search")
        web_search = st.checkbox("Enable web search")
        
        st.subheader("💬 Chat Mode")
        chat_mode = st.radio("Select chat mode", ["General Chat", "Document Chat"])
        
        if st.button("Clear Chat History"):
            st.session_state.messages = []
            st.rerun()

    # Main content
    if uploaded_files:
        qa_chain = process_documents(uploaded_files)
        if qa_chain:
            st.success("Documents processed successfully!")
        else:
            st.warning("No valid documents were uploaded or processed.")
    else:
        qa_chain = None

    # Chat interface
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if prompt := st.chat_input("What is your question?"):
        st.chat_message("user").markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})

        if chat_mode == "General Chat" or not qa_chain:
            response, _, web_results = get_chatbot_response(prompt, personality=personality, web_search=web_search)
        else:
            response, source_docs, web_results = get_chatbot_response(prompt, qa_chain, personality, web_search)
        
        with st.chat_message("assistant"):
            st.markdown(response)
            if chat_mode == "Document Chat" and qa_chain and source_docs:
                with st.expander("Source Documents"):
                    for doc in source_docs:
                        st.markdown(f"**Source:** {doc.metadata.get('source', 'Unknown')}")
                        st.markdown(doc.page_content[:200] + "...")
            
            if web_search and web_results:
                with st.expander("Web Search Results"):
                    for result in web_results:
                        st.markdown(f"**[{result['title']}]({result['link']})**")
                        st.markdown(result['snippet'])
        
        st.session_state.messages.append({"role": "assistant", "content": response})

    # Chat history and download
    with st.sidebar:
        st.subheader("📜 Chat History")
        history_expander = st.expander("View Chat History")
        with history_expander:
            for message in st.session_state.messages:
                st.text(f"{message['role']}: {message['content'][:50]}...")
        
        if st.session_state.messages:
            chat_history_df = pd.DataFrame(st.session_state.messages)
            csv = chat_history_df.to_csv(index=False)
            st.download_button(
                label="📥 Download Chat History",
                data=csv,
                file_name="chat_history.csv",
                mime="text/csv",
            )

    st.sidebar.markdown("---")
    st.sidebar.markdown("Powered by Falcon-180B and Streamlit")

if __name__ == "__main__":
    main()